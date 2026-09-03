import json
import re
from typing import Dict, Any, List
import pandas as pd

from schemas import NormalizedTRFSchema, DeviceInfoSchema, ExtractedStandardSchema, ExtractedTestSchema

class DocumentProcessor:
    @staticmethod
    def process_file(file_path: str, file_type: str) -> NormalizedTRFSchema:
        ext = file_type.lower().strip(".")
        if ext == "pdf":
            return DocumentProcessor._process_pdf(file_path)
        elif ext == "csv":
            return DocumentProcessor._process_csv(file_path)
        elif ext in ["xlsx", "xls"]:
            return DocumentProcessor._process_xlsx(file_path)
        elif ext in ["docx", "doc"]:
            return DocumentProcessor._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: .{ext}")

    @staticmethod
    def _process_pdf(file_path: str) -> NormalizedTRFSchema:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz
        doc = fitz.open(file_path)
        full_text = ""
        extracted_tests: List[ExtractedTestSchema] = []

        for page in doc:
            full_text += page.get_text() + "\n"
            # Extract PDF tables using PyMuPDF table finder
            tabs = page.find_tables()
            if tabs and tabs.tables:
                for tab in tabs.tables:
                    rows = tab.extract()
                    if not rows:
                        continue
                    header = [str(c).strip().lower() for c in rows[0] if c]
                    is_test_table = any("test" in h or "parameter" in h or "result" in h or "observation" in h for h in header)
                    start_idx = 1 if is_test_table else 0

                    for row in rows[start_idx:]:
                        if not row or len(row) < 2:
                            continue
                        test_name = str(row[0]).strip() if row[0] else ""
                        if not test_name or test_name.lower() in ["test parameter", "parameter", "test name", "test"]:
                            continue

                        result_raw = str(row[1]).strip() if len(row) > 1 and row[1] is not None else ""
                        unit_raw = str(row[2]).strip() if len(row) > 2 and row[2] is not None else ""
                        evidence_raw = str(row[3]).strip() if len(row) > 3 and row[3] is not None else ""

                        try:
                            clean_res = float(result_raw)
                        except (ValueError, TypeError):
                            clean_res = result_raw if result_raw else None

                        if unit_raw in ["-", "N/A", "None", ""]:
                            unit_raw = None

                        extracted_tests.append(ExtractedTestSchema(
                            test_name=test_name,
                            result=clean_res,
                            unit=unit_raw,
                            evidence=evidence_raw
                        ))

        doc.close()

        # Parse device info and standards from text
        norm = DocumentProcessor._parse_extracted_text(full_text)

        # Merge extracted table tests
        if extracted_tests:
            existing_names = {t.test_name.lower() for t in norm.tests}
            for et in extracted_tests:
                if et.test_name.lower() not in existing_names:
                    norm.tests.append(et)
                    existing_names.add(et.test_name.lower())

        return norm

    @staticmethod
    def _process_csv(file_path: str) -> NormalizedTRFSchema:
        df = pd.read_csv(file_path)
        return DocumentProcessor._parse_dataframe(df)

    @staticmethod
    def _process_xlsx(file_path: str) -> NormalizedTRFSchema:
        df = pd.read_excel(file_path)
        return DocumentProcessor._parse_dataframe(df)

    @staticmethod
    def _process_docx(file_path: str) -> NormalizedTRFSchema:
        import docx
        doc = docx.Document(file_path)
        full_text = []
        for p in doc.paragraphs:
            if p.text:
                full_text.append(p.text)
        
        # Also parse tables in docx
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                full_text.append(row_text)

        text_str = "\n".join(full_text)
        return DocumentProcessor._parse_extracted_text(text_str)

    @staticmethod
    def _parse_dataframe(df: pd.DataFrame) -> NormalizedTRFSchema:
        # Standardize column headers
        df.columns = [str(c).strip().lower() for c in df.columns]

        device_name = "Demo Medical Device"
        model = "BP-100"
        manufacturer = "Demo Med Tech Inc"
        device_category = "Blood Pressure Monitor"

        tests = []
        for _, row in df.iterrows():
            # Look for test name, value, unit
            test_name = row.get("test_name") or row.get("test") or row.get("parameter")
            if not test_name or pd.isna(test_name):
                continue
            
            result_val = row.get("result") or row.get("value") or row.get("measured_value") or row.get("observed_value")
            unit_val = row.get("unit") or row.get("units")
            evidence_val = row.get("evidence") or row.get("notes") or row.get("technician")
            
            # Handle float conversions if possible
            if pd.isna(result_val):
                clean_res = None
            else:
                try:
                    clean_res = float(result_val)
                except (ValueError, TypeError):
                    clean_res = str(result_val).strip()

            tests.append(ExtractedTestSchema(
                test_name=str(test_name).strip(),
                result=clean_res,
                unit=None if pd.isna(unit_val) else str(unit_val).strip(),
                evidence=None if pd.isna(evidence_val) else str(evidence_val).strip()
            ))

        return NormalizedTRFSchema(
            device=DeviceInfoSchema(
                name=device_name,
                model=model,
                manufacturer=manufacturer,
                device_type=device_category
            ),
            standards=[ExtractedStandardSchema(name="IEC 60601-1", edition="Demo Edition")],
            tests=tests,
            raw_notes="Parsed from spreadsheet data."
        )

    @staticmethod
    def _parse_extracted_text(text: str) -> NormalizedTRFSchema:
        # Extract Device Meta via Regex fallbacks
        device_name_m = re.search(r"Device(?:\s+Name)?:\s*[\r\n]*([^\r\n]+)", text, re.IGNORECASE)
        model_m = re.search(r"Model(?:\s+Number)?:\s*[\r\n]*([^\r\n]+)", text, re.IGNORECASE)
        mfr_m = re.search(r"Manufacturer:\s*[\r\n]*([^\r\n]+)", text, re.IGNORECASE)
        type_m = re.search(r"Device\s+Category:\s*[\r\n]*([^\r\n]+)", text, re.IGNORECASE)
        if not type_m:
            type_m = re.search(r"Device\s+Type:\s*[\r\n]*([^\r\n]+)", text, re.IGNORECASE)
        date_m = re.search(r"Test\s+Date:\s*[\r\n]*([^\r\n]+)", text, re.IGNORECASE)
        serial_m = re.search(r"Serial(?:\s+Number)?:\s*[\r\n]*([^\r\n]+)", text, re.IGNORECASE)

        device_info = DeviceInfoSchema(
            name=device_name_m.group(1).strip() if device_name_m else "Demo Blood Pressure Monitor",
            model=model_m.group(1).strip() if model_m else "BP-100",
            manufacturer=mfr_m.group(1).strip() if mfr_m else "Demo Health Medical Systems",
            device_type=type_m.group(1).strip() if type_m else "Blood Pressure Monitor",
            serial_number=serial_m.group(1).strip() if serial_m else "SN-994821",
            test_date=date_m.group(1).strip() if date_m else "2026-09-01"
        )

        # Extract Standards
        standards = []
        std_matches = re.findall(r"(IEC\s*\d+(?:-\d+)?|DEMO-STD-[A-Z]+)", text, re.IGNORECASE)
        for s in set(std_matches):
            standards.append(ExtractedStandardSchema(name=s.upper(), edition="Demo"))
        if not standards:
            standards.append(ExtractedStandardSchema(name="IEC 60601-1", edition="Demo"))

        # Extract Tests via patterns
        tests: List[ExtractedTestSchema] = []
        
        # Pattern line matchers for common TRF lines like "Leakage Current: 0.22 mA | PASS"
        lines = text.split("\n")
        for line in lines:
            line_str = line.strip()
            if not line_str or ":" not in line_str:
                continue

            # Look for lines with test name and values
            parts = line_str.split(":", 1)
            candidate_test = parts[0].strip()
            rest = parts[1].strip()

            # Skip header lines
            if candidate_test.lower() in ["device name", "model", "manufacturer", "device type", "serial number", "test date", "standard"]:
                continue

            # Try extracting number and unit from rest
            val_match = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*([a-zA-Z°Ω%]+)?", rest)
            if val_match:
                try:
                    num_val = float(val_match.group(1))
                    unit_val = val_match.group(2) if val_match.group(2) else None
                    tests.append(ExtractedTestSchema(
                        test_name=candidate_test,
                        result=num_val,
                        unit=unit_val,
                        evidence=rest
                    ))
                    continue
                except ValueError:
                    pass
            
            # Check categorical values like PASS / FAIL / COMPLETE
            cat_match = re.search(r"(PASS|FAIL|COMPLETE|SATISFACTORY)", rest, re.IGNORECASE)
            if cat_match:
                tests.append(ExtractedTestSchema(
                    test_name=candidate_test,
                    result=cat_match.group(1).upper(),
                    unit=None,
                    evidence=rest
                ))
                continue

            # If line seems like a test but has no clear result
            if len(candidate_test) > 3 and ("test" in candidate_test.lower() or "verification" in candidate_test.lower() or "resistance" in candidate_test.lower()):
                tests.append(ExtractedTestSchema(
                    test_name=candidate_test,
                    result=rest if rest else None,
                    unit=None,
                    evidence=rest
                ))

        return NormalizedTRFSchema(
            device=device_info,
            standards=standards,
            tests=tests,
            raw_notes=text[:1000]
        )
